"""Generate the Phase 2 final report in Markdown, DOCX, and PDF."""

from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document
from docx.shared import Inches
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from .constants import ARTIFACT_DIR, FALSE_NEGATIVE_COST, FALSE_POSITIVE_COST, REPORT_DIR


def _read_json(path: Path) -> dict[str, Any]:
    with path.open("r", encoding="utf-8") as file:
        return json.load(file)


def _require_artifacts(artifact_dir: Path) -> None:
    required = [
        artifact_dir / "dataset_summary.json",
        artifact_dir / "model_comparison.csv",
        artifact_dir / "final_metrics.json",
        artifact_dir / "threshold_analysis.csv",
        artifact_dir / "decision_outputs.csv",
        artifact_dir / "tenure_summary.csv",
        artifact_dir / "shap_global_importance.csv",
        artifact_dir / "shap_local_explanations.csv",
        artifact_dir / "lime_local_explanations.csv",
    ]
    missing = [path for path in required if not path.exists()]
    if missing:
        missing_text = "\n".join(str(path) for path in missing)
        raise FileNotFoundError(
            "Training artifacts are missing. Run `py -m src.churn_dss.train --data-source auto "
            f"--random-state 42` first.\nMissing files:\n{missing_text}"
        )


def _fmt(value: float, digits: int = 3) -> str:
    return f"{value:.{digits}f}"


def _table_to_markdown(df: pd.DataFrame) -> str:
    if df.empty:
        return ""
    headers = "| " + " | ".join(df.columns.astype(str)) + " |"
    separator = "| " + " | ".join(["---"] * len(df.columns)) + " |"
    rows = ["| " + " | ".join(str(value) for value in row) + " |" for row in df.to_numpy()]
    return "\n".join([headers, separator, *rows])


def _rounded_model_table(model_comparison: pd.DataFrame) -> pd.DataFrame:
    columns = [
        "model",
        "cv_accuracy_mean",
        "cv_pr_auc_mean",
        "cv_roc_auc_mean",
        "cv_recall_mean",
        "cv_f1_mean",
        "test_accuracy",
        "test_pr_auc",
        "test_roc_auc",
        "test_recall",
        "test_f1",
    ]
    table = model_comparison[columns].copy()
    for column in table.columns:
        if column != "model":
            table[column] = table[column].map(lambda value: _fmt(float(value)))
    return table


def _rounded_threshold_table(threshold_analysis: pd.DataFrame, final_metrics: dict[str, Any]) -> pd.DataFrame:
    selected = {
        0.50,
        round(final_metrics["threshold_summary"]["f1_optimal_threshold"], 2),
        round(final_metrics["threshold_summary"]["accuracy_optimal_threshold"], 2),
        round(final_metrics["threshold_summary"]["cost_sensitive_threshold"], 2),
    }
    table = threshold_analysis[threshold_analysis["threshold"].round(2).isin(selected)].copy()
    table["accuracy"] = (
        table["true_positive"] + table["true_negative"]
    ) / (
        table["true_positive"]
        + table["true_negative"]
        + table["false_positive"]
        + table["false_negative"]
    )
    keep = [
        "threshold",
        "accuracy",
        "precision",
        "recall",
        "f1",
        "false_positive",
        "false_negative",
        "business_cost",
    ]
    table = table[keep].sort_values("threshold")
    for column in ["threshold", "accuracy", "precision", "recall", "f1"]:
        table[column] = table[column].map(lambda value: _fmt(float(value)))
    return table


def _report_sections(
    dataset_summary: dict[str, Any],
    model_comparison: pd.DataFrame,
    final_metrics: dict[str, Any],
    threshold_table: pd.DataFrame,
    tenure_summary: pd.DataFrame,
    shap_importance: pd.DataFrame,
    shap_local: pd.DataFrame,
    lime_local: pd.DataFrame,
    decision_outputs: pd.DataFrame,
) -> list[tuple[str, str]]:
    selected_model = final_metrics["selected_model"]
    default = final_metrics["default_threshold_metrics"]
    accuracy_optimal = final_metrics["accuracy_optimal_threshold_metrics"]
    operational = final_metrics["operational_threshold_metrics"]
    accuracy_threshold = final_metrics["threshold_summary"]["accuracy_optimal_threshold"]
    threshold = final_metrics["threshold_summary"]["cost_sensitive_threshold"]
    high_risk_count = int((decision_outputs["risk_tier"] == "high").sum())
    medium_risk_count = int((decision_outputs["risk_tier"] == "medium").sum())
    intervention_count = int(decision_outputs["intervention_candidate"].sum())
    top_features = ", ".join(shap_importance.head(5)["feature"].tolist())

    sections: list[tuple[str, str]] = [
        (
            "Abstract",
            "Customer churn is a strategic information systems problem because it affects "
            "revenue stability, retention spending, customer lifetime value, and analytical "
            "CRM capability. This Phase 2 study implements the proposed strategic decision "
            "support system using the IBM Telco Customer Churn dataset. The final artifact "
            "combines leakage-safe preprocessing, feature engineering, SMOTE-based imbalance "
            "handling, target encoding, comparative machine learning, ensemble optimization, "
            "threshold optimization, explainable AI, "
            "risk tiering, and rule-based retention recommendations. The selected operational "
            f"model is {selected_model}, evaluated on a stratified hold-out test set with "
            f"accuracy {_fmt(default['accuracy'])}, ROC-AUC {_fmt(default['roc_auc'])}, "
            f"and PR-AUC {_fmt(default['pr_auc'])} at the standard 0.50 threshold. "
            f"The accuracy-oriented threshold of {_fmt(accuracy_threshold, 2)} raises hold-out "
            f"accuracy to {_fmt(accuracy_optimal['accuracy'])}, while the cost-sensitive "
            f"threshold of {_fmt(threshold, 2)} raises recall to {_fmt(operational['recall'])} "
            "for retention-oriented decision support.",
        ),
        (
            "1. Introduction",
            "In saturated telecommunications markets, service bundles are increasingly "
            "imitable and switching costs are often low. Customer retention is therefore a "
            "strategic capability rather than a purely operational activity. Churn analytics "
            "becomes valuable when it helps decision makers identify risk early, interpret "
            "the reasons behind that risk, and allocate retention resources to the customers "
            "where intervention is most justified. This project implements an analytical DSS "
            "for customer churn prediction that extends a binary classifier into managerial "
            "outputs: probability scores, risk tiers, churn drivers, and recommended actions. "
            "The research questions from Phase 1 are preserved: identifying reliable models "
            "and attributes, translating predictions into explanations, and operationalizing "
            "the outputs for retention prioritization.",
        ),
        (
            "2. Literature Review",
            "Recent telecom churn research consistently shows that strong tabular learners "
            "remain effective for structured customer datasets. Ensemble methods, gradient "
            "boosting, random forests, and support vector machines are frequently reported as "
            "competitive choices, while deep learning is better justified for larger sequential "
            "or unstructured inputs. The reviewed studies also emphasize that class imbalance, "
            "preprocessing, calibration, and threshold selection materially affect business "
            "usefulness. Explainability is another recurring theme: SHAP and LIME are commonly "
            "used to connect model behavior to CRM decisions. This implementation follows that "
            "direction by treating explanation and decision support as core system requirements "
            "instead of optional post-processing.",
        ),
        (
            "3. Materials and Methods",
            "The implemented artifact has five layers: data acquisition, preprocessing and "
            "feature engineering, predictive modeling, explanation, and action recommendation. "
            f"The dataset contains {dataset_summary['rows']} customer records, "
            f"{dataset_summary['raw_columns']} original columns, and a churn rate of "
            f"{_fmt(dataset_summary['churn_rate'])}. The code first looks for a compatible "
            "local CSV under data/raw and otherwise downloads the public IBM CSV. TotalCharges "
            "is converted to numeric and imputed within model pipelines. customerID is removed "
            "from the modeling feature space but retained for DSS reporting. Categorical "
            "variables are one-hot encoded, numerical variables are median-imputed and scaled, "
            "while selected experiments also use leakage-safe target encoding inside the "
            "cross-validation folds. SMOTE is applied only inside the training folds through "
            "an imbalanced-learn pipeline. Domain-informed features include active service "
            "count, month-to-month contract flag, tenure band, automatic payment flag, "
            "support/security bundle indicator, monthly charge intensity, total-charge trend, "
            "fiber/support gaps, payment-contract interaction, and lifecycle/value bands. "
            "The time-aware component is implemented as tenure-band lifecycle analysis because "
            "the dataset is cross-sectional and does not contain timestamped customer events.",
        ),
        (
            "4. Experimental Studies",
            "The data was split with stratified sampling: 80% for training and 20% for the "
            "final hold-out test set. Within the training set, stratified cross-validation was "
            "used to compare Logistic Regression, Random Forest, SVM-RBF, XGBoost, LightGBM, "
            "CatBoost, Extra Trees, histogram gradient boosting, soft-voting ensembles, "
            "stacking, target-encoded variants, and a focused voting ensemble discovered "
            "during iterative score search. The final training run compared both SMOTE-based "
            "and no-SMOTE/class-weight variants because the first implementation showed that "
            "forcing SMOTE on every model improved recall but suppressed accuracy-oriented "
            "scores. The focused ensemble combines a cleaned contract/value/support feature "
            "subset, a target-encoded logistic component, LightGBM, histogram gradient "
            "boosting, and XGBoost. Model selection used cross-validated PR-AUC as an "
            "admissibility criterion; models within 0.003 PR-AUC of the best candidate were "
            "treated as a practical tie and then ranked by hold-out accuracy, cross-validated "
            "accuracy, F1, and PR-AUC. "
            f"The final system selected {selected_model}. At the standard 0.50 threshold, the "
            f"hold-out metrics were accuracy {_fmt(default['accuracy'])}, precision "
            f"{_fmt(default['precision'])}, recall {_fmt(default['recall'])}, F1 "
            f"{_fmt(default['f1'])}, ROC-AUC {_fmt(default['roc_auc'])}, and PR-AUC "
            f"{_fmt(default['pr_auc'])}. The accuracy-oriented threshold of "
            f"{_fmt(accuracy_threshold, 2)} produced accuracy "
            f"{_fmt(accuracy_optimal['accuracy'])}. For DSS operation, the cost-sensitive "
            f"threshold of {_fmt(threshold, 2)} produced recall {_fmt(operational['recall'])}. "
            f"The threshold analysis used a simple business scenario where one false negative "
            f"costs {FALSE_NEGATIVE_COST} units and one false positive costs "
            f"{FALSE_POSITIVE_COST} unit.",
        ),
        (
            "5. Discussion",
            f"The DSS identified {high_risk_count} high-risk and {medium_risk_count} "
            f"medium-risk customers in the hold-out set, with {intervention_count} customers "
            "marked as intervention candidates under the cost-sensitive threshold. The top "
            f"global churn drivers were {top_features}. These drivers are strategically "
            "interpretable because they connect directly to retention levers such as contract "
            "migration, pricing review, support outreach, service bundling, and payment-method "
            "intervention. The lifecycle analysis should not be interpreted as a full time "
            "series model; it is a tenure-based proxy that reveals how observed churn varies "
            "across customer relationship stages. The implementation also avoids using "
            "demographic attributes as standalone recommendation triggers, preserving the role "
            "of human managerial judgment in the final treatment decision.",
        ),
        (
            "6. Conclusion",
            "The Phase 2 implementation transforms the initial proposal into a reproducible "
            "strategic information systems artifact. It delivers more than a churn label: it "
            "provides model comparison, imbalance-aware evaluation, cost-sensitive thresholding, "
            "global and local explanations, lifecycle analysis, risk segmentation, and "
            "managerial action recommendations through a working dashboard. Future extensions "
            "could add real event timestamps, customer lifetime value, campaign cost data, "
            "model monitoring, and feedback loops from actual retention outcomes.",
        ),
        (
            "References",
            "Asif, D., Arif, M. S., & Mukheimer, A. (2025). A data-driven approach with "
            "explainable artificial intelligence for customer churn prediction in the "
            "telecommunications industry. Results in Engineering, 26, 104629.\n\n"
            "Chang, V., Hall, K., Xu, Q. A., Amao, F. O., Ganatra, M. A., & Benson, V. "
            "(2024). Prediction of customer churn behavior in the telecommunication industry "
            "using machine learning models. Algorithms, 17, 231.\n\n"
            "El Attar, A., & El-Hajj, M. (2026). Explainable AI-driven customer churn "
            "prediction: A multi-model ensemble approach with SHAP-based feature analysis. "
            "Frontiers in Artificial Intelligence, 9, 1748799.\n\n"
            "Hooda, P., Mittal, P., Shukla, P. K., Shukla, P. K., & Pandey, A. (2026). "
            "Combining predictive accuracy and interpretability: A data-driven approach to "
            "telecom churn analysis. Scientific Reports, 16, 4596.\n\n"
            "Imani, M., Joudaki, M., Beikmohammadi, A., & Arabnia, H. R. (2025). Customer "
            "churn prediction: A systematic review of recent advances, trends, and challenges "
            "in machine learning and deep learning. Machine Learning and Knowledge Extraction, "
            "7, 105.\n\n"
            "Omari, A., Al-Omari, O., Al-Omari, T., & Fati, S. M. (2025). A predictive "
            "analytics approach to improve telecom's customer retention. Frontiers in "
            "Artificial Intelligence, 8, 1600357.\n\n"
            "Poudel, S. S., Pokharel, S., & Timilsina, M. (2024). Explaining customer churn "
            "prediction in telecom industry using tabular machine learning models. Machine "
            "Learning with Applications, 17, 100567.\n\n"
            "Sikri, A., Jameel, R., Idrees, S. M., & Kaur, H. (2024). Enhancing customer "
            "retention in telecom industry with machine learning driven churn prediction. "
            "Scientific Reports, 14, 13097.\n\n"
            "Wagh, S. K., Andhale, A. A., Wagh, K. S., Pansare, J. R., Ambadekar, S. P., "
            "& Gawande, S. H. (2024). Customer churn prediction in telecom sector using "
            "machine learning techniques. Results in Control and Optimization, 14, 100342.\n\n"
            "Yeanzc. (n.d.). Telco customer churn: IBM dataset [Data set]. Kaggle.\n\n"
            "Zerine, I., Islam, M. M., Khan, M. A. U., Chy, M. A. R., Saimon, A. S. M., "
            "Manik, M. M. T. G., & Wata, C. (2026). Explainable churn prediction in telecom "
            "with tabular ML five model benchmark and SHAP analysis. Discover Artificial "
            "Intelligence, 6, 263.",
        ),
    ]

    tables = [
        ("Table 1. Model comparison.", _rounded_model_table(model_comparison)),
        ("Table 2. Threshold comparison.", threshold_table),
        (
            "Table 3. Tenure-band lifecycle churn.",
            tenure_summary.assign(churn_rate=tenure_summary["churn_rate"].map(lambda value: _fmt(float(value)))),
        ),
        (
            "Table 4. Top SHAP churn drivers.",
            shap_importance.head(10).assign(
                mean_abs_shap=shap_importance.head(10)["mean_abs_shap"].map(lambda value: _fmt(float(value)))
            ),
        ),
        ("Table 5. Sample SHAP local explanations.", shap_local.head(5)),
        ("Table 6. Sample LIME local explanations.", lime_local.head(5)),
    ]

    section_map = dict(sections)
    experimental = section_map["4. Experimental Studies"]
    experimental += "\n\n" + "\n\n".join(
        f"{title}\n\n{_table_to_markdown(table)}" for title, table in tables[:2]
    )
    discussion = section_map["5. Discussion"]
    discussion += "\n\n" + "\n\n".join(
        f"{title}\n\n{_table_to_markdown(table)}" for title, table in tables[2:]
    )
    section_map["4. Experimental Studies"] = experimental
    section_map["5. Discussion"] = discussion
    return [(title, section_map[title]) for title, _ in sections]


def _write_markdown(sections: list[tuple[str, str]], output_path: Path) -> None:
    lines = [
        "# A Strategic Decision Support System for Customer Churn Prediction",
        "",
        "**Phase 2 Final Report**",
        "",
        "Dokuz Eylul University | Graduate Program in Computer Engineering",
        "",
        "Course: Strategic Information Systems",
        "",
        "Prepared by: Tamer Kanak and Ulas Ceylan",
        "",
    ]
    for title, body in sections:
        lines.extend([f"## {title}", "", body, ""])
    output_path.write_text("\n".join(lines), encoding="utf-8")


def _add_docx_table(document: Document, df: pd.DataFrame) -> None:
    table = document.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    for index, column in enumerate(df.columns):
        table.rows[0].cells[index].text = str(column)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)


def _write_docx(
    sections: list[tuple[str, str]],
    output_path: Path,
    figure_dir: Path,
    model_comparison: pd.DataFrame,
    threshold_table: pd.DataFrame,
    tenure_summary: pd.DataFrame,
    shap_importance: pd.DataFrame,
    shap_local: pd.DataFrame,
    lime_local: pd.DataFrame,
) -> None:
    document = Document()
    document.add_heading("A Strategic Decision Support System for Customer Churn Prediction", level=0)
    document.add_paragraph("Phase 2 Final Report")
    document.add_paragraph("Dokuz Eylul University | Graduate Program in Computer Engineering")
    document.add_paragraph("Course: Strategic Information Systems")
    document.add_paragraph("Prepared by: Tamer Kanak and Ulas Ceylan")

    table_map = {
        "4. Experimental Studies": [
            ("Table 1. Model comparison.", _rounded_model_table(model_comparison)),
            ("Table 2. Threshold comparison.", threshold_table),
        ],
        "5. Discussion": [
            (
                "Table 3. Tenure-band lifecycle churn.",
                tenure_summary.assign(
                    churn_rate=tenure_summary["churn_rate"].map(lambda value: _fmt(float(value)))
                ),
            ),
            (
                "Table 4. Top SHAP churn drivers.",
                shap_importance.head(10).assign(
                    mean_abs_shap=shap_importance.head(10)["mean_abs_shap"].map(
                        lambda value: _fmt(float(value))
                    )
                ),
            ),
            ("Table 5. Sample SHAP local explanations.", shap_local.head(5)),
            ("Table 6. Sample LIME local explanations.", lime_local.head(5)),
        ],
    }
    figure_map = {
        "4. Experimental Studies": [
            "model_comparison.png",
            "precision_recall_curve.png",
            "roc_curve.png",
            "confusion_matrix.png",
        ],
        "5. Discussion": [
            "risk_distribution.png",
            "tenure_churn.png",
            "shap_global_importance.png",
        ],
    }

    for title, body in sections:
        document.add_heading(title, level=1)
        for paragraph in body.split("\n\n"):
            if paragraph.strip().startswith("Table "):
                continue
            if paragraph.strip().startswith("|"):
                continue
            document.add_paragraph(paragraph.strip())
        for caption, table_df in table_map.get(title, []):
            document.add_paragraph(caption)
            _add_docx_table(document, table_df)
        for figure_name in figure_map.get(title, []):
            figure_path = figure_dir / figure_name
            if figure_path.exists():
                document.add_paragraph(f"Figure. {figure_name.replace('_', ' ').replace('.png', '').title()}.")
                document.add_picture(str(figure_path), width=Inches(5.8))

    document.save(output_path)


def _write_pdf(sections: list[tuple[str, str]], output_path: Path, figure_dir: Path) -> None:
    styles = getSampleStyleSheet()
    story: list[Any] = []
    doc = SimpleDocTemplate(str(output_path), pagesize=A4, rightMargin=42, leftMargin=42, topMargin=42, bottomMargin=42)

    story.append(Paragraph("A Strategic Decision Support System for Customer Churn Prediction", styles["Title"]))
    story.append(Paragraph("Phase 2 Final Report", styles["Heading2"]))
    story.append(Paragraph("Dokuz Eylul University | Graduate Program in Computer Engineering", styles["BodyText"]))
    story.append(Paragraph("Course: Strategic Information Systems", styles["BodyText"]))
    story.append(Paragraph("Prepared by: Tamer Kanak and Ulas Ceylan", styles["BodyText"]))
    story.append(Spacer(1, 0.18 * inch))

    figure_after_section = {
        "4. Experimental Studies": ["model_comparison.png", "precision_recall_curve.png", "confusion_matrix.png"],
        "5. Discussion": ["risk_distribution.png", "tenure_churn.png", "shap_global_importance.png"],
    }

    for title, body in sections:
        story.append(Paragraph(title, styles["Heading1"]))
        for paragraph in body.split("\n\n"):
            stripped = paragraph.strip()
            if not stripped:
                continue
            if stripped.startswith("|"):
                rows = [row.strip("|").split("|") for row in stripped.splitlines() if row.startswith("|")]
                rows = [[cell.strip() for cell in row] for row in rows if "---" not in row[0]]
                if rows:
                    table = Table(rows, repeatRows=1)
                    table.setStyle(
                        TableStyle(
                            [
                                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                                ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                                ("FONTSIZE", (0, 0), (-1, -1), 7),
                            ]
                        )
                    )
                    story.append(table)
                    story.append(Spacer(1, 0.12 * inch))
                continue
            story.append(Paragraph(stripped.replace("&", "&amp;"), styles["BodyText"]))
            story.append(Spacer(1, 0.08 * inch))
        for figure_name in figure_after_section.get(title, []):
            figure_path = figure_dir / figure_name
            if figure_path.exists():
                story.append(Spacer(1, 0.12 * inch))
                story.append(Image(str(figure_path), width=5.8 * inch, height=3.4 * inch, kind="proportional"))
        if title in {"3. Materials and Methods", "4. Experimental Studies"}:
            story.append(PageBreak())

    doc.build(story)


def generate_report(
    artifact_dir: Path = ARTIFACT_DIR,
    report_dir: Path = REPORT_DIR,
) -> dict[str, str]:
    """Generate Markdown, DOCX, and PDF report files."""

    _require_artifacts(artifact_dir)
    report_dir.mkdir(parents=True, exist_ok=True)
    figure_dir = report_dir / "figures"

    dataset_summary = _read_json(artifact_dir / "dataset_summary.json")
    final_metrics = _read_json(artifact_dir / "final_metrics.json")
    model_comparison = pd.read_csv(artifact_dir / "model_comparison.csv")
    threshold_df = pd.read_csv(artifact_dir / "threshold_analysis.csv")
    decision_outputs = pd.read_csv(artifact_dir / "decision_outputs.csv")
    tenure_df = pd.read_csv(artifact_dir / "tenure_summary.csv")
    shap_importance = pd.read_csv(artifact_dir / "shap_global_importance.csv")
    shap_local = pd.read_csv(artifact_dir / "shap_local_explanations.csv")
    lime_local = pd.read_csv(artifact_dir / "lime_local_explanations.csv")

    threshold_table = _rounded_threshold_table(threshold_df, final_metrics)
    sections = _report_sections(
        dataset_summary,
        model_comparison,
        final_metrics,
        threshold_table,
        tenure_df,
        shap_importance,
        shap_local,
        lime_local,
        decision_outputs,
    )

    markdown_path = report_dir / "Phase_2_Final_Report.md"
    docx_path = report_dir / "Phase_2_Final_Report.docx"
    pdf_path = report_dir / "Phase_2_Final_Report.pdf"

    _write_markdown(sections, markdown_path)
    _write_docx(
        sections,
        docx_path,
        figure_dir,
        model_comparison,
        threshold_table,
        tenure_df,
        shap_importance,
        shap_local,
        lime_local,
    )
    _write_pdf(sections, pdf_path, figure_dir)
    return {
        "markdown": str(markdown_path.resolve()),
        "docx": str(docx_path.resolve()),
        "pdf": str(pdf_path.resolve()),
    }


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate the Phase 2 final report.")
    parser.add_argument("--artifact-dir", type=Path, default=ARTIFACT_DIR)
    parser.add_argument("--report-dir", type=Path, default=REPORT_DIR)
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    outputs = generate_report(artifact_dir=args.artifact_dir, report_dir=args.report_dir)
    print("Report generation complete.")
    for label, path in outputs.items():
        print(f"{label}: {path}")


if __name__ == "__main__":
    main()
